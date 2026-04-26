"""
server.py — Company Catcher DART 뷰어 서버
─────────────────────────────────────────
DART 공시 데이터베이스 조회 및 웹 뷰어 제공

API 엔드포인트:
  GET  /api/dart/stats                — 수집 현황 통계
  GET  /api/dart/companies            — 기업 목록 (검색/필터/페이지)
  GET  /api/dart/reports              — 특정 기업 보고서 목록
  GET  /api/dart/report               — 단일 보고서 전문 (raw_text + biz_content)
  GET  /api/dart/supply_chain         — 공급망 단순 조회 (하위호환)
  GET  /api/dart/latest_type          — 가장 최근 수집된 보고서 유형
  GET  /api/ai/usage                  — AI 모델 일일 사용량 조회
  GET  /api/session                   — 세션 토큰 (로컬호스트 전용)
  POST /api/ai/analyze                — AI 비교 분석 (X-Api-Key 필요)

  공급망 그래프 API:
  GET  /api/supply-chain/stats        — 공급망 DB 전체 통계
  GET  /api/supply-chain/graph        — 기업 중심 그래프 (?corp_code&depth&direction&relation)
  GET  /api/supply-chain/hub          — 파트너명 역방향 조회 (?partner_name&relation)
  GET  /api/supply-chain/impact       — 충격 전파 분석 (?corp_code)
  GET  /api/supply-chain/path         — 두 기업 간 연결 경로 (?from&to&max_depth)
  GET  /api/supply-chain/hubs         — 허브 순위 (?sector&limit)
  GET  /api/supply-chain/changes      — 보고서 기간별 공급망 변화 (?corp_code)
  GET  /api/supply-chain/common       — 교집합 종목 발굴 (?partners=A&partners=B&relation)
  GET  /api/supply-chain/concentration — 집중도 리스크 스코어 (?corp_code)
  GET  /api/supply-chain/risk-ranking  — 집중도 리스크 순위 (?limit&risk_level)
  GET  /api/supply-chain/theme-article — 테마 기사 데이터 (?corp_code&type)
  GET  /api/supply-chain/risk-scenario — 리스크 시나리오 시뮬레이션 (?corp_code&scenario)

사용법:
  python server.py              → http://localhost:8888
  python server.py --port 9000
"""

import argparse
import json
import io
import logging
import os
import re
import secrets
import sqlite3
import sys
import threading
import time
import urllib.request
import urllib.error
from datetime import datetime
from http.server import BaseHTTPRequestHandler, HTTPServer
from pathlib import Path
from urllib.parse import parse_qs, urlparse, quote
import socketserver

# ── 선택적 내보내기 라이브러리 (없어도 서버 구동) ─────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    from fpdf import FPDF
    _HAS_FPDF = True
except ImportError:
    _HAS_FPDF = False

# Windows cp949 콘솔 인코딩 오류 방지
if sys.stdout.encoding and sys.stdout.encoding.lower() not in ("utf-8", "utf8"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

# ── 경로 설정 ──────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent
DB_PATH = ROOT / "data" / "dart" / "dart_reports.db"
WEB_ROOT = ROOT / "web"
AI_USAGE_PATH = ROOT / "data" / "dart" / "ai_usage.json"

# ── 로깅 ───────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s │ %(levelname)-7s │ %(message)s",
    datefmt="%H:%M:%S",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler(ROOT / "logs" / "server.log", encoding="utf-8"),
    ],
)
log = logging.getLogger("dart-server")

# ── 환경 변수 로드 ──────────────────────────────────────────────────────────
_env_path = ROOT / ".env"
if _env_path.exists():
    for line in _env_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            k, v = line.split("=", 1)
            key, val = k.strip(), v.strip()
            # .env 값이 있으면 OS 환경변수가 비어있거나 없을 때 덮어쓴다
            if val and not os.environ.get(key):
                os.environ[key] = val


def _ensure_api_secret_key():
    """API_SECRET_KEY 가 없으면 자동 생성하여 .env 에 저장"""
    if os.environ.get("API_SECRET_KEY"):
        return
    key = secrets.token_urlsafe(32)
    os.environ["API_SECRET_KEY"] = key
    # .env 파일에 추가
    existing = _env_path.read_text(encoding="utf-8") if _env_path.exists() else ""
    if "API_SECRET_KEY=" in existing:
        lines = []
        for ln in existing.splitlines():
            if ln.strip().startswith("API_SECRET_KEY=") or ln.strip().startswith("# API_SECRET_KEY="):
                lines.append(f"API_SECRET_KEY={key}")
            else:
                lines.append(ln)
        _env_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    else:
        with open(_env_path, "a", encoding="utf-8") as f:
            f.write(f"\nAPI_SECRET_KEY={key}\n")
    log.info(f"API_SECRET_KEY 자동 생성됨")


# ── DB 연결 (thread-local) ─────────────────────────────────────────────────
_local = threading.local()
_db_lock = threading.Lock()


def get_db() -> sqlite3.Connection:
    """스레드 로컬 DB 연결 반환"""
    if not hasattr(_local, "conn") or _local.conn is None:
        conn = sqlite3.connect(str(DB_PATH), check_same_thread=False)
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA synchronous=NORMAL")
        _local.conn = conn
    return _local.conn


# ── AI 사용량 관리 ─────────────────────────────────────────────────────────
_usage_lock = threading.Lock()


def load_ai_usage() -> dict:
    today = datetime.now().strftime("%Y-%m-%d")
    if AI_USAGE_PATH.exists():
        try:
            data = json.loads(AI_USAGE_PATH.read_text(encoding="utf-8"))
            if data.get("date") == today:
                return data
        except Exception:
            pass
    return {"date": today, "flash_used": 0, "pro_used": 0, "sonnet_used": 0}


def save_ai_usage(data: dict):
    AI_USAGE_PATH.write_text(
        json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def increment_usage(model: str):
    with _usage_lock:
        usage = load_ai_usage()
        key = f"{model}_used"
        usage[key] = usage.get(key, 0) + 1
        save_ai_usage(usage)


# ── AI 호출 함수 ────────────────────────────────────────────────────────────
# 무료 API 일일 한도 (2026 기준)
GEMINI_LIMITS = {
    "flash":      {"rpm": 10, "rpd": 500},
    "flash-lite": {"rpm": 15, "rpd": 1000},
    "pro":        {"rpm":  5, "rpd": 100},
    "article":    {"rpm": 10, "rpd": 100},   # Gemini 3.1 Pro Preview (기사 초안 전용)
}

# 모델별 fallback 순서 (404 시 자동 시도)
GEMINI_FALLBACKS = {
    "flash":      ["gemini-2.5-flash-preview-05-20",
                   "gemini-2.5-flash",
                   "gemini-2.0-flash"],
    "flash-lite": ["gemini-2.5-flash-lite-preview-06-17",
                   "gemini-2.5-flash-lite",
                   "gemini-2.5-flash-8b",
                   "gemini-2.0-flash-lite"],
    "pro":        ["gemini-2.5-pro-preview-05-06",
                   "gemini-2.5-pro",
                   "gemini-2.0-pro-exp"],
    "article":    ["gemini-3.1-pro-preview",
                   "gemini-2.5-pro-preview-05-06",
                   "gemini-2.5-pro"],
}


def _get_gemini_keys() -> list:
    """파싱1 + 파싱2 API 키 목록 반환"""
    keys = []
    k1 = os.environ.get("GEMINI_API_KEY", "").strip()
    k2 = os.environ.get("GEMINI_API_KEY_2", "").strip()
    if k1: keys.append(k1)
    if k2: keys.append(k2)
    return keys

_srv_key_idx = 0

def call_gemini(model_type: str, prompt: str) -> str:
    global _srv_key_idx
    keys = _get_gemini_keys()
    if not keys:
        raise ValueError("GEMINI_API_KEY가 .env에 설정되지 않았습니다")

    candidates = GEMINI_FALLBACKS.get(model_type, GEMINI_FALLBACKS["flash"])
    last_err = None

    for model_name in candidates:
        # 라운드로빈으로 키 선택
        api_key = keys[_srv_key_idx % len(keys)]
        _srv_key_idx += 1

        url = (
            f"https://generativelanguage.googleapis.com/v1beta/models/"
            f"{model_name}:generateContent?key={api_key}"
        )
        payload = json.dumps({
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"maxOutputTokens": 8192, "temperature": 0.3},
        }).encode("utf-8")

        req = urllib.request.Request(
            url, data=payload,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(req, timeout=120) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            key_label = f"파싱{keys.index(api_key)+1}"
            log.info(f"Gemini 사용 모델: {model_name} ({key_label})")
            return data["candidates"][0]["content"]["parts"][0]["text"]
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="replace")
            if e.code == 404:
                log.warning(f"Gemini 모델 없음({model_name}), 다음 시도...")
                last_err = f"모델 없음: {model_name}"
                continue
            if e.code == 429 and len(keys) > 1:
                # 다른 키로 재시도
                alt_key = keys[(keys.index(api_key) + 1) % len(keys)]
                alt_url = url.replace(api_key, alt_key)
                alt_req = urllib.request.Request(alt_url, data=payload,
                    headers={"Content-Type": "application/json"}, method="POST")
                try:
                    with urllib.request.urlopen(alt_req, timeout=120) as r2:
                        d2 = json.loads(r2.read().decode("utf-8"))
                    return d2["candidates"][0]["content"]["parts"][0]["text"]
                except Exception:
                    pass
            raise RuntimeError(f"Gemini API 오류 {e.code}: {body[:400]}")

    raise RuntimeError(f"사용 가능한 Gemini 모델 없음. 마지막 오류: {last_err}")


def call_claude(prompt: str) -> str:
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY가 .env에 설정되지 않았습니다")

    url = "https://api.anthropic.com/v1/messages"
    payload = json.dumps({
        "model": "claude-sonnet-4-6",
        "max_tokens": 8192,
        "messages": [{"role": "user", "content": prompt}],
    }).encode("utf-8")

    req = urllib.request.Request(
        url, data=payload,
        headers={
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        return data["content"][0]["text"]
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"Claude API 오류 {e.code}: {body[:400]}")


def build_compare_prompt(reports: list, focus: str = "") -> str:
    parts = []
    for i, r in enumerate(reports, 1):
        parts.append(f"[{r['label']}]\n{r['content']}")

    focus_line = f"\n특별히 '{focus}' 관련 내용을 중점 분석해주세요.\n" if focus else ""

    sep = "─" * 60
    combined = f"\n{sep}\n".join(parts)

    return f"""아래는 동일 기업의 DART 공시 사업보고서 중 **'2. 사업의 내용'** 섹션 {len(reports)}개입니다.
각 보고서의 사업 내용만을 기반으로 비교 분석해주세요.{focus_line}

{sep}
{combined}
{sep}

## 분석 요청 항목

1. **핵심 변화 요약**
   - 기간별로 사업 내용에서 달라진 핵심 사항을 표로 정리

2. **사업 전략 변화**
   - 신규 사업 진출, 기존 사업 확장·축소·철수
   - 주력 제품·서비스의 변화

3. **시장 및 경쟁환경 변화**
   - 주요 고객, 시장 점유율, 경쟁사 관련 서술 변화

4. **리스크 요인 변화**
   - 새롭게 등장하거나 해소된 리스크 (원재료, 규제, 기술, 시장)

5. **수치·실적 언급 변화**
   - 사업 내용에 등장하는 매출·물량·점유율 등 구체 수치 비교

6. **투자자 관점 인사이트**
   - 사업 내용에서 포착할 수 있는 중요 시그널 (긍정/부정)

분석은 한국어로 작성하고, 각 항목마다 보고서 원문의 구체적인 표현·수치를 근거로 제시해주세요."""


# ── MIME 타입 ──────────────────────────────────────────────────────────────
MIME = {
    ".html": "text/html; charset=utf-8",
    ".css":  "text/css",
    ".js":   "application/javascript",
    ".json": "application/json",
    ".svg":  "image/svg+xml",
    ".png":  "image/png",
    ".ico":  "image/x-icon",
}

# ── 기사 초안 생성 헬퍼 (인라인) ──────────────────────────────────────────────
_ARTICLE_MODELS = [
    "gemini-3.1-pro-preview",
    "gemini-2.5-pro-preview-05-06",
    "gemini-2.5-pro-preview-06-05",
    "gemini-2.5-pro",
]
_LEAD_TYPE_KO = {
    "strategy_change": "경영전략 변화",
    "market_shift":    "시장 변화",
    "risk_alert":      "리스크 경보",
    "numeric_change":  "수치 급변",
    "supply_chain":    "공급망 변화",
}
_ARTICLE_STYLE_GUIDE = """[기사 작성 지침]
- 매체: 파이낸스코프 (finscope.co.kr) / 기자: 고종민
- 문체: 연합뉴스·중앙일보 경제면 스타일 (간결, 객관, 사실 중심)
- 제목: 30자 이내, 핵심 사실 중심 (과장 금지)
- 부제: 15자 이내
- 본문: 400~600자, 역피라미드 구조 (리드→구체내용→의미→전망)
- 주의: DART 공시 내용만 근거, 추측 시 '~것으로 분석된다' 표현"""


def _build_article_prompt(lead, ai_result: str) -> str:
    corp_name = lead["corp_name"] or "해당 기업"
    type_ko   = _LEAD_TYPE_KO.get(lead["lead_type"], lead["lead_type"])
    try:
        kw_list = ", ".join(json.loads(lead["keywords"] or "[]"))
    except Exception:
        kw_list = lead["keywords"] or ""
    return f"""당신은 파이낸스코프(finscope.co.kr)의 경제 전문 기자 고종민입니다.
아래 DART 공시 분석 데이터를 바탕으로 한국어 경제 기사 초안을 작성하세요.

{_ARTICLE_STYLE_GUIDE}

═══════════════════════════════
[취재 단서]
기업명: {corp_name}
단서 유형: {type_ko} (severity {lead['severity']}/5)
감지 키워드: {kw_list}
단서 제목: {lead['title'] or ''}
핵심 요약: {lead['summary'] or ''}
근거 문장: {lead['evidence'] or ''}

[AI 분석 참고 (최대 2500자)]
{(ai_result or '')[:2500]}
═══════════════════════════════

반드시 아래 JSON 형식으로만 출력하세요:
{{
  "headline": "기사 제목 (30자 이내)",
  "subheadline": "부제 (15자 이내)",
  "body": "본문 (400~600자, 단락 사이 빈 줄)",
  "keywords": ["키워드1", "키워드2", "키워드3"],
  "news_value": "취재 가치 한 줄 설명",
  "caution": "확인 필요 사항 (없으면 빈 문자열)"
}}"""


def _call_gemini_article(prompt: str, timeout: int = 120) -> tuple:
    global _srv_key_idx
    keys = _get_gemini_keys()
    if not keys:
        raise ValueError("GEMINI_API_KEY 없음")
    # 라운드로빈으로 키 선택
    api_key = keys[_srv_key_idx % len(keys)]
    _srv_key_idx += 1
    last_err = None
    for model_name in _ARTICLE_MODELS:
        url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
               f"{model_name}:generateContent?key={api_key}")
        payload = json.dumps({
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"maxOutputTokens": 4096, "temperature": 0.7},
        }).encode("utf-8")
        req = urllib.request.Request(url, data=payload,
            headers={"Content-Type": "application/json"}, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            return data["candidates"][0]["content"]["parts"][0]["text"], model_name
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8", errors="replace")
            if e.code == 404:
                last_err = f"모델 없음: {model_name}"; continue
            if e.code == 429:
                raise RuntimeError(f"API 할당량 초과(429) — 내일 다시 시도하세요.")
            raise RuntimeError(f"Gemini API 오류 {e.code}: {body[:300]}")
        except urllib.error.URLError as e:
            raise RuntimeError(f"네트워크 오류: {e.reason}")
    raise RuntimeError(f"사용 가능한 모델 없음. 마지막 오류: {last_err}")


def _parse_article_json(text: str) -> dict:
    m = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if m: text = m.group(1)
    else:
        m2 = re.search(r"(\{.*\})", text, re.DOTALL)
        if m2: text = m2.group(1)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        result = {}
        for field in ["headline", "subheadline", "body", "news_value", "caution"]:
            fm = re.search(rf'"{field}"\s*:\s*"(.*?)"(?=\s*[,\}}])', text, re.DOTALL)
            if fm: result[field] = fm.group(1).replace("\\n", "\n").strip()
        km = re.search(r'"keywords"\s*:\s*\[(.*?)\]', text, re.DOTALL)
        if km: result["keywords"] = re.findall(r'"([^"]+)"', km.group(1))
        return result


def _save_article_draft(conn, lead_id: int, lead, article: dict, model_name: str) -> int:
    headline    = (article.get("headline") or "")[:200]
    subheadline = (article.get("subheadline") or "")[:200]
    body        = article.get("body") or ""
    news_value  = article.get("news_value") or ""
    caution     = article.get("caution") or ""
    editor_note = ""
    if news_value: editor_note += f"[취재 가치] {news_value}\n"
    if caution:    editor_note += f"[주의] {caution}"
    cur = conn.execute("""
        INSERT INTO article_drafts
            (lead_id, corp_code, corp_name,
             headline, subheadline, content,
             style, model, word_count, char_count,
             status, editor_note, created_at)
        VALUES (?,?,?,?,?,?,?,?,?,?,'draft',?,datetime('now','localtime'))
    """, (lead_id, lead["corp_code"], lead["corp_name"],
          headline, subheadline, body,
          "news", model_name,
          len(body.split()), len(body), editor_note.strip()))
    conn.commit()
    return cur.lastrowid


# ── HTTP 핸들러 ────────────────────────────────────────────────────────────
class DartHandler(BaseHTTPRequestHandler):

    def log_message(self, fmt, *args):
        pass  # BaseHTTP 기본 로그 억제

    def _send(self, code: int, body: bytes, ct: str = "application/json"):
        self.send_response(code)
        self.send_header("Content-Type", ct)
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()
        self.wfile.write(body)

    def _json(self, data, code: int = 200):
        body = json.dumps(data, ensure_ascii=False, indent=2, default=str).encode("utf-8")
        self._send(code, body)

    def _raw(self, buf: bytes, ct: str, filename: str):
        """바이너리 파일 다운로드 응답"""
        self.send_response(200)
        self.send_header("Content-Type", ct)
        self.send_header("Content-Length", str(len(buf)))
        self.send_header("Access-Control-Allow-Origin", "*")
        safe = filename.encode("ascii", "ignore").decode()
        enc  = quote(filename)
        self.send_header("Content-Disposition",
                         f"attachment; filename=\"{safe}\"; filename*=UTF-8''{enc}")
        self.end_headers()
        self.wfile.write(buf)

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type, X-Api-Key")
        self.end_headers()

    def _check_api_key(self) -> bool:
        expected = os.environ.get("API_SECRET_KEY", "")
        provided = self.headers.get("X-Api-Key", "")
        return not expected or secrets.compare_digest(expected, provided)

    def do_GET(self):
        parsed = urlparse(self.path)
        path   = parsed.path
        qs     = parse_qs(parsed.query)

        # ── SPA 페이지 라우트 (확장자 없는 경로 + 상세 페이지) ──────────────
        _spa = {
            "/leads":              "leads.html",
            "/comparisons":        "comparisons.html",
            "/articles":           "articles.html",
            "/lead_detail":        "lead_detail.html",
            "/article_detail":     "article_detail.html",
            "/comparison_detail":  "comparison_detail.html",
            "/supply_chain":       "supply_chain.html",
            "/supply_chain_leads": "supply_chain_leads.html",
            "/company_detail":     "company_detail.html",
            "/alert_rules":        "alert_rules.html",
            "/home_v2":            "home_v2.html",
            "/dashboard":          "home_v2.html",
        }
        if path in _spa:
            f = WEB_ROOT / _spa[path]
            if f.is_file():
                self._send(200, f.read_bytes(), "text/html; charset=utf-8")
            else:
                self._send(404, b"Not Found", "text/plain")
            return

        # ── /api/session (로컬호스트 전용 — API 키 반환) ─────────────────
        if path == "/api/session":
            client_ip = self.client_address[0]
            if client_ip not in ("127.0.0.1", "::1"):
                self._json({"error": "로컬에서만 접근 가능"}, 403)
                return
            self._json({
                "api_key": os.environ.get("API_SECRET_KEY", ""),
                "has_gemini": bool(os.environ.get("GEMINI_API_KEY")),
                "has_claude": bool(os.environ.get("ANTHROPIC_API_KEY")),
            })
            return

        # ── /api/dart/stats ───────────────────────────────────────────────
        if path == "/api/dart/stats":
            try:
                db = get_db()
                total       = db.execute("SELECT COUNT(*) FROM reports").fetchone()[0]
                with_content = db.execute(
                    "SELECT COUNT(*) FROM reports WHERE "
                    "(raw_text IS NOT NULL AND raw_text != '') OR "
                    "(biz_content IS NOT NULL AND biz_content != '')"
                ).fetchone()[0]
                companies   = db.execute(
                    "SELECT COUNT(DISTINCT corp_code) FROM reports WHERE rcept_no != '' AND "
                    "((raw_text IS NOT NULL AND raw_text != '') OR "
                    " (biz_content IS NOT NULL AND biz_content != ''))"
                ).fetchone()[0]
                supply_chain = db.execute("SELECT COUNT(*) FROM supply_chain").fetchone()[0]
                annual  = db.execute("SELECT COUNT(*) FROM reports WHERE report_type='2025_annual'").fetchone()[0]
                q3      = db.execute("SELECT COUNT(*) FROM reports WHERE report_type='2025_q3'").fetchone()[0]
                h1      = db.execute("SELECT COUNT(*) FROM reports WHERE report_type='2025_h1'").fetchone()[0]
                q1      = db.execute("SELECT COUNT(*) FROM reports WHERE report_type='2025_q1'").fetchone()[0]
                old_a   = db.execute("SELECT COUNT(*) FROM reports WHERE report_type='A'").fetchone()[0]
                self._json({
                    "total": total,
                    "with_text": with_content,
                    "companies": companies,
                    "supply_chain": supply_chain,
                    "annual_2025": annual,
                    "q3_2025": q3,
                    "h1_2025": h1,
                    "q1_2025": q1,
                    "legacy_A": old_a,
                })
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/report_types ─────────────────────────────────────────────
        # 소프트코딩: DB의 report_type_meta + 실제 수집 현황 합산 반환
        if path == "/api/report_types":
            try:
                db = get_db()
                # report_type_meta 등록 유형
                meta_rows = db.execute("""
                    SELECT * FROM report_type_meta ORDER BY sort_order
                """).fetchall()
                # 실제 DB에 수집된 유형별 건수
                counts = {r["report_type"]: r["cnt"] for r in db.execute("""
                    SELECT report_type,
                           COUNT(*) as cnt
                    FROM reports
                    WHERE biz_content IS NOT NULL AND LENGTH(biz_content) > 100
                    GROUP BY report_type
                """).fetchall()}
                # 합산
                result = []
                for m in meta_rows:
                    tc = m["type_code"]
                    result.append({
                        **dict(m),
                        "count": counts.get(tc, 0),
                        "has_data": counts.get(tc, 0) > 0,
                    })
                # 메타에 없지만 실제 수집된 유형 추가 (레거시 등)
                meta_codes = {m["type_code"] for m in meta_rows}
                for tc, cnt in counts.items():
                    if tc not in meta_codes:
                        result.append({
                            "type_code": tc, "label": tc, "short_label": tc,
                            "year": None, "quarter": None,
                            "count": cnt, "has_data": True, "is_active": 1,
                        })
                self._json({"report_types": result})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/dart/latest_type ─────────────────────────────────────────
        if path == "/api/dart/latest_type":
            try:
                db = get_db()
                # 가장 최근 rcept_dt를 기준으로 보고서 유형 결정 (동적)
                row = db.execute("""
                    SELECT report_type, MAX(rcept_dt) as latest_dt
                    FROM reports
                    WHERE rcept_no != ''
                      AND biz_content IS NOT NULL AND LENGTH(biz_content) > 100
                    GROUP BY report_type
                    ORDER BY latest_dt DESC
                    LIMIT 1
                """).fetchone()
                if row:
                    self._json({"report_type": row["report_type"], "rcept_dt": row["latest_dt"]})
                else:
                    self._json({"report_type": None})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/ai/usage ─────────────────────────────────────────────────
        if path == "/api/ai/usage":
            try:
                usage = load_ai_usage()
                self._json({
                    "date":   usage["date"],
                    "flash":      {"used": usage.get("flash_used",      0), "limit": 500},
                    "flash-lite": {"used": usage.get("flash_lite_used", 0), "limit": 1000},
                    "pro":        {"used": usage.get("pro_used",        0), "limit": 100},
                    "sonnet":     {"used": usage.get("sonnet_used",     0), "limit": None},
                })
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/dart/companies ───────────────────────────────────────────
        if path == "/api/dart/companies":
            try:
                db = get_db()
                q      = (qs.get("q",    [None])[0] or "").strip()
                rtype  = (qs.get("type", [None])[0] or "").strip()
                limit  = int(qs.get("limit",  [100])[0])
                offset = int(qs.get("offset", [0])[0])

                where = [
                    "rcept_no != ''",
                    "((raw_text IS NOT NULL AND raw_text != '') OR "
                    " (biz_content IS NOT NULL AND biz_content != ''))",
                ]
                params = []

                if q:
                    where.append("(corp_name LIKE ? OR corp_code LIKE ?)")
                    params += [f"%{q}%", f"%{q}%"]

                if rtype:
                    where.append("report_type = ?")
                    params.append(rtype)

                where_sql = " AND ".join(where)

                count_sql = f"SELECT COUNT(DISTINCT corp_code) FROM reports WHERE {where_sql}"
                total = db.execute(count_sql, params).fetchone()[0]

                sql = f"""
                    SELECT
                        corp_code,
                        MAX(corp_name) as corp_name,
                        COUNT(*) as report_cnt,
                        SUM(CASE WHEN report_type='2025_annual' THEN 1 ELSE 0 END) as annual_cnt,
                        SUM(CASE WHEN report_type='2025_q3'     THEN 1 ELSE 0 END) as q3_cnt,
                        SUM(CASE WHEN report_type='2025_h1'     THEN 1 ELSE 0 END) as h1_cnt,
                        SUM(CASE WHEN report_type='2025_q1'     THEN 1 ELSE 0 END) as q1_cnt,
                        SUM(CASE WHEN report_type='A'           THEN 1 ELSE 0 END) as legacy_cnt,
                        MAX(rcept_dt) as last_rcept_dt
                    FROM reports
                    WHERE {where_sql}
                    GROUP BY corp_code
                    ORDER BY
                        CASE WHEN MAX(corp_name) IS NULL OR MAX(corp_name) = '' THEN 1 ELSE 0 END,
                        MAX(corp_name) COLLATE NOCASE
                    LIMIT ? OFFSET ?
                """
                rows = db.execute(sql, params + [limit, offset]).fetchall()
                companies = [dict(r) for r in rows]
                self._json({"companies": companies, "total": total, "offset": offset, "limit": limit})
            except Exception as e:
                log.exception("companies error")
                self._json({"error": str(e)}, 500)
            return

        # ── /api/dart/reports ─────────────────────────────────────────────
        if path == "/api/dart/reports":
            try:
                db = get_db()
                corp_code = (qs.get("corp_code", [None])[0] or "").strip()
                rtype     = (qs.get("type",      [None])[0] or "").strip()
                limit     = int(qs.get("limit",  [50])[0])
                offset    = int(qs.get("offset", [0])[0])

                if not corp_code:
                    self._json({"error": "corp_code 필수"}, 400)
                    return

                has_content = (
                    "((raw_text IS NOT NULL AND raw_text != '') OR "
                    " (biz_content IS NOT NULL AND biz_content != ''))"
                )
                where = [f"corp_code = ?", has_content]
                params = [corp_code]
                if rtype:
                    where.append("report_type = ?")
                    params.append(rtype)

                where_sql = " AND ".join(where)
                rows = db.execute(f"""
                    SELECT
                        id, corp_code, corp_name, report_type, report_name,
                        rcept_no, rcept_dt, flr_nm,
                        LENGTH(raw_text)    as text_len,
                        LENGTH(biz_content) as biz_len,
                        CASE WHEN biz_content IS NOT NULL AND LENGTH(biz_content) >= 300
                             THEN 1 ELSE 0 END as has_biz,
                        SUBSTR(COALESCE(biz_content, raw_text, ''), 1, 300) as preview
                    FROM reports
                    WHERE {where_sql}
                    ORDER BY rcept_dt DESC
                    LIMIT ? OFFSET ?
                """, params + [limit, offset]).fetchall()

                reports = [dict(r) for r in rows]
                self._json({"reports": reports, "corp_code": corp_code})
            except Exception as e:
                log.exception("reports error")
                self._json({"error": str(e)}, 500)
            return

        # ── /api/dart/report ──────────────────────────────────────────────
        if path == "/api/dart/report":
            try:
                db = get_db()
                rcept_no = (qs.get("rcept_no", [None])[0] or "").strip()
                rec_id   = (qs.get("id",       [None])[0] or "").strip()

                if rcept_no:
                    row = db.execute(
                        "SELECT * FROM reports WHERE rcept_no = ?", [rcept_no]
                    ).fetchone()
                elif rec_id:
                    row = db.execute(
                        "SELECT * FROM reports WHERE id = ?", [int(rec_id)]
                    ).fetchone()
                else:
                    self._json({"error": "rcept_no 또는 id 필수"}, 400)
                    return

                if not row:
                    self._json({"error": "보고서 없음"}, 404)
                    return
                self._json(dict(row))
            except Exception as e:
                log.exception("report error")
                self._json({"error": str(e)}, 500)
            return

        # ── /api/settings/status ─────────────────────────────────────────
        if path == "/api/settings/status":
            client_ip = self.client_address[0]
            if client_ip not in ("127.0.0.1", "::1"):
                self._json({"error": "로컬에서만 접근 가능"}, 403)
                return
            gemini_key  = os.environ.get("GEMINI_API_KEY",    "")
            gemini_key2 = os.environ.get("GEMINI_API_KEY_2",  "")
            claude_key  = os.environ.get("ANTHROPIC_API_KEY", "")
            self._json({
                "has_gemini":   bool(gemini_key),
                "has_gemini2":  bool(gemini_key2),
                "has_claude":   bool(claude_key),
                "gemini_hint":  (gemini_key[:8]  + "…") if gemini_key  else "",
                "gemini2_hint": (gemini_key2[:8] + "…") if gemini_key2 else "",
                "claude_hint":  (claude_key[:10] + "…") if claude_key  else "",
            })
            return

        # ── /api/search ───────────────────────────────────────────────────
        # FTS5 전문 검색: 기업명 + 사업내용 키워드
        if path == "/api/search":
            try:
                db = get_db()
                q       = (qs.get("q",      [None])[0] or "").strip()
                rtype   = (qs.get("type",   [None])[0] or "").strip()
                sector  = (qs.get("sector", [None])[0] or "").strip()
                limit   = int(qs.get("limit",  [20])[0])
                offset  = int(qs.get("offset", [0])[0])

                if not q:
                    self._json({"error": "q(검색어) 필수"}, 400)
                    return

                # FTS5 쿼리 — 기업명 OR 사업내용
                fts_q = q.replace('"', '""')
                where_extra = ""
                params_extra = []
                if rtype:
                    where_extra += " AND r.report_type = ?"
                    params_extra.append(rtype)

                rows = db.execute(f"""
                    SELECT r.id, r.corp_code, r.corp_name, r.report_type,
                           r.report_name, r.rcept_dt,
                           LENGTH(r.biz_content) as biz_len,
                           snippet(reports_fts, 2, '<mark>', '</mark>', '...', 32) as snippet
                    FROM reports_fts f
                    JOIN reports r ON r.id = f.rowid
                    WHERE reports_fts MATCH ?
                      {where_extra}
                    ORDER BY rank
                    LIMIT ? OFFSET ?
                """, [f'"{fts_q}"'] + params_extra + [limit, offset]).fetchall()

                # 전체 수 (근사)
                total_row = db.execute(f"""
                    SELECT COUNT(*) FROM reports_fts f
                    JOIN reports r ON r.id = f.rowid
                    WHERE reports_fts MATCH ?
                    {where_extra}
                """, [f'"{fts_q}"'] + params_extra).fetchone()
                total = total_row[0] if total_row else 0

                # 연관 비교분석·취재단서 수 추가 (기업코드 기준)
                corp_codes = list({r["corp_code"] for r in rows if r["corp_code"]})
                cmp_counts, lead_counts = {}, {}
                if corp_codes:
                    ph = ",".join("?" * len(corp_codes))
                    for row in db.execute(
                        f"SELECT corp_code, COUNT(*) FROM ai_comparisons WHERE corp_code IN ({ph}) AND status='ok' GROUP BY corp_code",
                        corp_codes
                    ).fetchall():
                        cmp_counts[row[0]] = row[1]
                    for row in db.execute(
                        f"SELECT corp_code, COUNT(*) FROM story_leads WHERE corp_code IN ({ph}) AND status!='archived' GROUP BY corp_code",
                        corp_codes
                    ).fetchall():
                        lead_counts[row[0]] = row[1]

                results = []
                for r in rows:
                    d2 = dict(r)
                    d2["cmp_count"]  = cmp_counts.get(r["corp_code"], 0)
                    d2["lead_count"] = lead_counts.get(r["corp_code"], 0)
                    results.append(d2)

                self._json({
                    "results": results,
                    "total": total,
                    "query": q,
                    "offset": offset,
                    "limit": limit,
                })
            except Exception as e:
                log.exception("search error")
                self._json({"error": str(e)}, 500)
            return

        # ── /api/alerts ────────────────────────────────────────────────────
        # 변화 알림 목록 (심각도순)
        if path == "/api/alerts":
            try:
                db = get_db()
                min_sev  = int(qs.get("min_severity", [3])[0])
                status   = (qs.get("status", ["new"])[0] or "new").strip()
                lead_type= (qs.get("lead_type",[None])[0] or "").strip()
                limit    = int(qs.get("limit",  [50])[0])
                offset   = int(qs.get("offset", [0])[0])

                where = ["severity >= ?", "status = ?"]
                params = [min_sev, status]
                if lead_type:
                    where.append("lead_type = ?")
                    params.append(lead_type)

                where_sql = " AND ".join(where)
                try:
                    total = db.execute(
                        f"SELECT COUNT(*) FROM story_leads WHERE {where_sql}", params
                    ).fetchone()[0]
                    rows = db.execute(
                        f"SELECT * FROM story_leads WHERE {where_sql} "
                        f"ORDER BY severity DESC, created_at DESC LIMIT ? OFFSET ?",
                        params + [limit, offset]
                    ).fetchall()
                    items = [dict(r) for r in rows]
                except Exception:
                    total, items = 0, []

                # 심각도별 통계
                try:
                    stats = db.execute("""
                        SELECT severity, lead_type, status, COUNT(*) as cnt
                        FROM story_leads
                        GROUP BY severity, lead_type, status
                        ORDER BY severity DESC
                    """).fetchall()
                    stats_list = [dict(r) for r in stats]
                except Exception:
                    stats_list = []

                self._json({
                    "alerts": items,
                    "total": total,
                    "stats": stats_list,
                    "offset": offset,
                    "limit": limit,
                })
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/leads ─────────────────────────────────────────────────────
        if path == "/api/leads":
            try:
                db = get_db()
                corp_code = (qs.get("corp_code",[None])[0] or "").strip()
                status    = (qs.get("status",  [None])[0] or "").strip()
                lead_type = (qs.get("type",    [None])[0] or "").strip()
                min_sev   = int(qs.get("min_severity",[1])[0])
                limit     = int(qs.get("limit", [30])[0])
                offset    = int(qs.get("offset",[0])[0])
                q         = (qs.get("q",       [None])[0] or "").strip()
                sort      = (qs.get("sort",    ["recent"])[0] or "recent").strip()

                where = ["severity >= ?"]
                params = [min_sev]
                if corp_code:
                    where.append("corp_code = ?"); params.append(corp_code)
                if status:
                    where.append("status = ?"); params.append(status)
                if lead_type:
                    where.append("lead_type = ?"); params.append(lead_type)
                if q:
                    where.append(
                        "(corp_name LIKE ? OR title LIKE ? OR summary LIKE ? OR evidence LIKE ?)"
                    )
                    params += [f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"]

                where_sql = " AND ".join(where)

                # 정렬
                if sort == "severity":
                    order_sql = "severity DESC, created_at DESC"
                elif sort == "relevant" and q:
                    order_sql = "severity DESC, created_at DESC"
                else:  # recent (default)
                    order_sql = "created_at DESC"

                try:
                    total = db.execute(
                        f"SELECT COUNT(*) FROM story_leads WHERE {where_sql}", params
                    ).fetchone()[0]
                    rows = db.execute(
                        f"""SELECT sl.*,
                              COALESCE(sl.draft_count,
                                (SELECT COUNT(*) FROM article_drafts ad WHERE ad.lead_id=sl.id)
                              ) AS draft_count_live
                            FROM story_leads sl
                            WHERE {where_sql}
                            ORDER BY {order_sql} LIMIT ? OFFSET ?""",
                        params + [limit, offset]
                    ).fetchall()
                    leads_list = []
                    for r in rows:
                        d = dict(r)
                        # draft_count_live 우선 사용
                        d["draft_count"] = d.pop("draft_count_live", 0) or d.get("draft_count", 0)
                        leads_list.append(d)
                    self._json({"leads": leads_list, "total": total})
                except Exception:
                    self._json({"leads": [], "total": 0})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/articles ──────────────────────────────────────────────────
        if path == "/api/articles":
            try:
                db = get_db()
                corp_code = (qs.get("corp_code",[None])[0] or "").strip()
                status    = (qs.get("status",  [None])[0] or "").strip()
                limit     = int(qs.get("limit", [20])[0])
                offset    = int(qs.get("offset",[0])[0])
                q         = (qs.get("q",       [None])[0] or "").strip()
                sort      = (qs.get("sort",    ["recent"])[0] or "recent").strip()

                where = []
                params = []
                if corp_code:
                    where.append("corp_code = ?"); params.append(corp_code)
                if status:
                    where.append("status = ?"); params.append(status)
                if q:
                    where.append(
                        "(corp_name LIKE ? OR headline LIKE ? OR subheadline LIKE ?)"
                    )
                    params += [f"%{q}%", f"%{q}%", f"%{q}%"]

                where_sql = ("WHERE " + " AND ".join(where)) if where else ""

                # 정렬 (NULL 처리: NULL은 마지막에 배치 — SQLite 호환)
                if sort == "verify_high":
                    order_sql = "(av.score_total IS NULL), av.score_total DESC, ad.created_at DESC"
                elif sort == "verify_low":
                    order_sql = "(av.score_total IS NULL), av.score_total ASC, ad.created_at DESC"
                elif sort == "char_count":
                    order_sql = "ad.char_count DESC"
                elif sort == "relevant" and q:
                    order_sql = "ad.created_at DESC"
                else:  # recent (default)
                    order_sql = "ad.created_at DESC"

                try:
                    total = db.execute(
                        f"SELECT COUNT(*) FROM article_drafts {where_sql}", params
                    ).fetchone()[0]
                    # 검증 점수 LEFT JOIN (테이블 없으면 NULL)
                    has_verification_table = db.execute(
                        "SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name='article_verification'"
                    ).fetchone()[0]
                    select_extra = ""
                    join_extra = ""
                    if has_verification_table:
                        select_extra = (
                            ", av.score_total as verify_score, av.flag as verify_flag, "
                            "av.score_numeric, av.score_direction, av.score_evidence, av.score_grounding"
                        )
                        join_extra = " LEFT JOIN article_verification av ON ad.id = av.article_id"

                    # WHERE 절 컬럼명에 ad. 프리픽스 (충돌 방지)
                    qual_where = where_sql
                    if qual_where:
                        for col in ("corp_code", "status", "corp_name", "headline", "subheadline"):
                            qual_where = re.sub(rf"\b{col}\b(\s*[=<>!]| LIKE)", f"ad.{col}\\1", qual_where)
                    rows = db.execute(
                        f"SELECT ad.id, ad.lead_id, ad.corp_code, ad.corp_name, ad.headline, ad.subheadline, "
                        f"ad.style, ad.model, ad.word_count, ad.char_count, ad.status, ad.created_at, "
                        f"SUBSTR(ad.content, 1, 300) as content_preview"
                        f"{select_extra} "
                        f"FROM article_drafts ad{join_extra} "
                        f"{qual_where} "
                        f"ORDER BY {order_sql} "
                        f"LIMIT ? OFFSET ?",
                        params + [limit, offset]
                    ).fetchall()
                    self._json({"articles": [dict(r) for r in rows], "total": total})
                except Exception as e:
                    self._json({"articles": [], "total": 0, "_error": str(e)})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/articles/{id} ────────────────────────────────────────────
        m_art = re.match(r"^/api/articles/(\d+)$", path)
        if m_art:
            art_id = int(m_art.group(1))
            try:
                db = get_db()
                row = db.execute("SELECT * FROM article_drafts WHERE id=?", [art_id]).fetchone()
                if not row:
                    self._json({"error": f"article_id={art_id} 없음"}, 404)
                else:
                    self._json({"article": dict(row)})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/articles/{id}/verification ───────────────────────────────
        m_art_verify = re.match(r"^/api/articles/(\d+)/verification$", path)
        if m_art_verify:
            art_id = int(m_art_verify.group(1))
            try:
                db = get_db()
                # 테이블 없을 수도 있음
                exists = db.execute(
                    "SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name='article_verification'"
                ).fetchone()[0]
                if not exists:
                    self._json({}); return
                row = db.execute(
                    "SELECT * FROM article_verification WHERE article_id=?", [art_id]
                ).fetchone()
                self._json(dict(row) if row else {})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/cross-signals (교차 신호) ────────────────────────────────
        if path == "/api/cross-signals":
            try:
                limit = int(qs.get("limit", ["30"])[0])
                pattern = qs.get("pattern", [None])[0]
                db = get_db()
                exists = db.execute(
                    "SELECT COUNT(*) FROM sqlite_master WHERE name='cross_signals'"
                ).fetchone()[0]
                if not exists:
                    self._json({"items": [], "total": 0}); return

                where = ""
                params = []
                if pattern:
                    where = "WHERE pattern = ?"
                    params.append(pattern.upper())
                total = db.execute(
                    f"SELECT COUNT(*) FROM cross_signals {where}", params
                ).fetchone()[0]
                rows = db.execute(f"""
                    SELECT * FROM cross_signals {where}
                    ORDER BY severity DESC, signal_count DESC, id DESC LIMIT ?
                """, params + [limit]).fetchall()
                self._json({"items": [dict(r) for r in rows], "total": total})
            except Exception as e:
                self._json({"items": [], "total": 0, "error": str(e)})
            return

        # ── /api/dashboard (홈 대시보드 통합) ─────────────────────────────
        if path == "/api/dashboard":
            try:
                db = get_db()
                result = {}

                # 1. 핫 단서 (severity 5 + 새로움 70+ 우선)
                try:
                    hot = db.execute("""
                        SELECT sl.id, sl.corp_name, sl.lead_type, sl.severity,
                               sl.title, sl.summary, sl.news_status,
                               ln.score_total as novelty,
                               ln.pattern as novelty_pattern,
                               cig.score_total as info_gap,
                               cig.grade as info_gap_grade,
                               (SELECT COUNT(*) FROM article_drafts ad
                                WHERE ad.lead_id = sl.id AND ad.char_count >= 500) as has_article
                        FROM story_leads sl
                        LEFT JOIN lead_novelty ln      ON sl.id = ln.lead_id
                        LEFT JOIN company_info_gap cig ON sl.corp_code = cig.corp_code
                        WHERE sl.severity >= 4
                        ORDER BY
                            COALESCE(ln.score_total, 50) DESC,
                            COALESCE(cig.score_total, 50) DESC,
                            sl.severity DESC
                        LIMIT 8
                    """).fetchall()
                    result["hot_leads"] = [dict(r) for r in hot]
                except Exception as e:
                    result["hot_leads"] = []
                    result["_hot_err"] = str(e)

                # 2. KPI 4개
                try:
                    kpi = {}
                    kpi["new_leads_24h"] = db.execute(
                        "SELECT COUNT(*) FROM story_leads WHERE created_at >= datetime('now', '-24 hours')"
                    ).fetchone()[0]
                    kpi["drafted_total"] = db.execute(
                        "SELECT COUNT(*) FROM article_drafts WHERE char_count >= 500"
                    ).fetchone()[0]
                    kpi["safe_articles"] = db.execute("""
                        SELECT COUNT(*) FROM article_verification WHERE flag IN ('SAFE','LOW')
                    """).fetchone()[0] if db.execute(
                        "SELECT COUNT(*) FROM sqlite_master WHERE name='article_verification'"
                    ).fetchone()[0] else 0
                    kpi["info_gap_high"] = db.execute("""
                        SELECT COUNT(*) FROM company_info_gap WHERE grade = '🚨'
                    """).fetchone()[0] if db.execute(
                        "SELECT COUNT(*) FROM sqlite_master WHERE name='company_info_gap'"
                    ).fetchone()[0] else 0
                    result["kpi"] = kpi
                except Exception as e:
                    result["kpi"] = {}
                    result["_kpi_err"] = str(e)

                # 3. 콘셉별 카운트
                try:
                    concepts = {}
                    concepts["story_leads_active"] = db.execute(
                        "SELECT COUNT(*) FROM story_leads WHERE severity >= 3"
                    ).fetchone()[0]
                    if db.execute("SELECT COUNT(*) FROM sqlite_master WHERE name='supply_chain_leads'").fetchone()[0]:
                        sc_rows = db.execute(
                            "SELECT scenario, COUNT(*) cnt FROM supply_chain_leads GROUP BY scenario"
                        ).fetchall()
                        concepts["supply_chain"] = {r["scenario"]: r["cnt"] for r in sc_rows}
                        concepts["supply_chain_total"] = sum(r["cnt"] for r in sc_rows)
                    if db.execute("SELECT COUNT(*) FROM sqlite_master WHERE name='event_disclosures'").fetchone()[0]:
                        concepts["disclosures_total"] = db.execute(
                            "SELECT COUNT(*) FROM event_disclosures"
                        ).fetchone()[0]
                    risk = db.execute(
                        "SELECT COUNT(*) FROM story_leads WHERE lead_type='risk_alert' AND severity >= 4"
                    ).fetchone()[0]
                    concepts["risk_alerts"] = risk
                    if db.execute("SELECT COUNT(*) FROM sqlite_master WHERE name='lead_novelty'").fetchone()[0]:
                        concepts["new_strong"] = db.execute(
                            "SELECT COUNT(*) FROM lead_novelty WHERE score_total >= 70"
                        ).fetchone()[0]
                    result["concepts"] = concepts
                except Exception as e:
                    result["concepts"] = {}
                    result["_concepts_err"] = str(e)

                # 4. 최근 수시공시 (severity 4+ 만)
                try:
                    if db.execute("SELECT COUNT(*) FROM sqlite_master WHERE name='event_leads'").fetchone()[0]:
                        rd = db.execute("""
                            SELECT el.corp_name, el.event_type, el.title, el.severity,
                                   ed.rcept_dt, ed.raw_url, ed.report_nm
                            FROM event_leads el
                            LEFT JOIN event_disclosures ed ON el.rcept_no = ed.rcept_no
                            WHERE el.severity >= 4
                            ORDER BY ed.rcept_dt DESC LIMIT 12
                        """).fetchall()
                        result["recent_disclosures"] = [dict(r) for r in rd]
                    else:
                        result["recent_disclosures"] = []
                except Exception as e:
                    result["recent_disclosures"] = []
                    result["_disc_err"] = str(e)

                # 4.5 교차 신호 (Phase Cross) — 사용자 사명 직결
                try:
                    if db.execute("SELECT COUNT(*) FROM sqlite_master WHERE name='cross_signals'").fetchone()[0]:
                        cs = db.execute("""
                            SELECT pattern, COUNT(*) cnt FROM cross_signals
                            GROUP BY pattern
                        """).fetchall()
                        cs_total = sum(r["cnt"] for r in cs)
                        cs_top = db.execute("""
                            SELECT pattern, title, interpretation, primary_corp_name,
                                   severity, signal_count
                            FROM cross_signals
                            ORDER BY severity DESC, signal_count DESC LIMIT 6
                        """).fetchall()
                        result["cross_signals"] = {
                            "total": cs_total,
                            "by_pattern": {r["pattern"]: r["cnt"] for r in cs},
                            "top": [dict(r) for r in cs_top],
                        }
                    else:
                        result["cross_signals"] = {"total": 0, "by_pattern": {}, "top": []}
                except Exception as e:
                    result["cross_signals"] = {"total": 0, "_err": str(e)}

                # 5. 최근 기사 (검증 점수 포함)
                try:
                    has_v = db.execute(
                        "SELECT COUNT(*) FROM sqlite_master WHERE name='article_verification'"
                    ).fetchone()[0]
                    if has_v:
                        la = db.execute("""
                            SELECT ad.id, ad.corp_name, ad.headline, ad.char_count,
                                   ad.created_at, ad.model,
                                   av.flag as verify_flag, av.score_total as verify_score,
                                   (SELECT COUNT(*) FROM article_fact_cards afc WHERE afc.article_id = ad.id) as is_v3
                            FROM article_drafts ad
                            LEFT JOIN article_verification av ON ad.id = av.article_id
                            WHERE ad.char_count >= 500
                            ORDER BY ad.id DESC LIMIT 6
                        """).fetchall()
                    else:
                        la = db.execute("""
                            SELECT id, corp_name, headline, char_count, created_at, model
                            FROM article_drafts WHERE char_count >= 500
                            ORDER BY id DESC LIMIT 6
                        """).fetchall()
                    result["latest_articles"] = [dict(r) for r in la]
                except Exception as e:
                    result["latest_articles"] = []
                    result["_la_err"] = str(e)

                self._json(result)
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain-leads (목록) ────────────────────────────────
        if path == "/api/supply-chain-leads":
            try:
                limit  = int(qs.get("limit",  ["30"])[0])
                offset = int(qs.get("offset", ["0"])[0])
                scenario = qs.get("scenario", [None])[0]
                db = get_db()
                exists = db.execute(
                    "SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name='supply_chain_leads'"
                ).fetchone()[0]
                if not exists:
                    self._json({"items": [], "total": 0}); return
                where = ""
                params = []
                if scenario:
                    where = "WHERE scenario = ?"
                    params.append(scenario)
                total = db.execute(
                    f"SELECT COUNT(*) FROM supply_chain_leads {where}", params
                ).fetchone()[0]
                rows = db.execute(f"""
                    SELECT * FROM supply_chain_leads {where}
                    ORDER BY severity DESC, id ASC LIMIT ? OFFSET ?
                """, params + [limit, offset]).fetchall()
                self._json({"items": [dict(r) for r in rows], "total": total})
            except Exception as e:
                self._json({"items": [], "total": 0, "error": str(e)})
            return

        # ── /api/articles/{id}/fact_card (v3) ─────────────────────────────
        m_art_fc = re.match(r"^/api/articles/(\d+)/fact_card$", path)
        if m_art_fc:
            art_id = int(m_art_fc.group(1))
            try:
                db = get_db()
                exists = db.execute(
                    "SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name='article_fact_cards'"
                ).fetchone()[0]
                if not exists:
                    self._json({}); return
                row = db.execute(
                    "SELECT * FROM article_fact_cards WHERE article_id=?", [art_id]
                ).fetchone()
                self._json(dict(row) if row else {})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/articles/{id}/reporter_brief (v3) ────────────────────────
        m_art_rb = re.match(r"^/api/articles/(\d+)/reporter_brief$", path)
        if m_art_rb:
            art_id = int(m_art_rb.group(1))
            try:
                db = get_db()
                exists = db.execute(
                    "SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name='reporter_briefs'"
                ).fetchone()[0]
                if not exists:
                    self._json({}); return
                row = db.execute(
                    "SELECT * FROM reporter_briefs WHERE article_id=?", [art_id]
                ).fetchone()
                self._json(dict(row) if row else {})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/articles/{id}/external_refs ──────────────────────────────
        m_art_ext = re.match(r"^/api/articles/(\d+)/external_refs$", path)
        if m_art_ext:
            art_id = int(m_art_ext.group(1))
            try:
                db = get_db()
                exists = db.execute(
                    "SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name='article_external_refs'"
                ).fetchone()[0]
                if not exists:
                    self._json({"refs": []}); return
                rows = db.execute("""
                    SELECT es.outlet_name, es.outlet_tier as tier,
                           es.title, es.summary, es.url, es.published_at,
                           aer.citation_text, aer.citation_position
                    FROM article_external_refs aer
                    JOIN external_sources es ON aer.source_id = es.id
                    WHERE aer.article_id = ?
                    ORDER BY aer.citation_position
                """, [art_id]).fetchall()
                self._json({"refs": [dict(r) for r in rows]})
            except Exception as e:
                self._json({"refs": [], "error": str(e)})
            return

        # ── /api/leads/{id}/external_pool ─────────────────────────────────
        m_lead_pool = re.match(r"^/api/leads/(\d+)/external_pool$", path)
        if m_lead_pool:
            lead_id = int(m_lead_pool.group(1))
            try:
                db = get_db()
                exists = db.execute(
                    "SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name='lead_external_match'"
                ).fetchone()[0]
                if not exists:
                    self._json({"items": []}); return
                rows = db.execute("""
                    SELECT es.outlet_name, es.outlet_tier as tier,
                           es.title, es.summary, es.url, es.published_at,
                           lem.match_score, lem.fact_check_status,
                           lem.fact_check_score
                    FROM lead_external_match lem
                    JOIN external_sources es ON lem.source_id = es.id
                    WHERE lem.lead_id = ?
                    ORDER BY
                      CASE lem.fact_check_status
                        WHEN 'PASS' THEN 0 WHEN 'PARTIAL' THEN 1
                        WHEN 'NOT_CHECKED' THEN 2 ELSE 3 END,
                      lem.match_score DESC
                    LIMIT 20
                """, [lead_id]).fetchall()
                self._json({"items": [dict(r) for r in rows]})
            except Exception as e:
                self._json({"items": [], "error": str(e)})
            return

        # ── /api/articles/{id}/export ─────────────────────────────────────
        m_art_export = re.match(r"^/api/articles/(\d+)/export$", path)
        if m_art_export:
            art_id = int(m_art_export.group(1))
            fmt    = qs.get("format", ["docx"])[0].lower()
            try:
                db  = get_db()
                row = db.execute("SELECT * FROM article_drafts WHERE id=?", [art_id]).fetchone()
                if not row:
                    self._json({"error": f"article_id={art_id} 없음"}, 404)
                    return
                a = dict(row)
                if fmt == "docx":
                    buf = _export_docx(a)
                    fname = _safe_fname(a.get("headline") or f"article_{art_id}") + ".docx"
                    self._raw(buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", fname)
                elif fmt == "pdf":
                    buf = _export_pdf(a)
                    fname = _safe_fname(a.get("headline") or f"article_{art_id}") + ".pdf"
                    self._raw(buf, "application/pdf", fname)
                else:
                    self._json({"error": "format must be docx or pdf"}, 400)
            except Exception as e:
                log.error("export error: %s", e, exc_info=True)
                self._json({"error": str(e)}, 500)
            return

        # ── /api/comparisons/{id}/export ─────────────────────────────────
        m_cmp_export = re.match(r"^/api/comparisons/(\d+)/export$", path)
        if m_cmp_export:
            cmp_id = int(m_cmp_export.group(1))
            fmt    = qs.get("format", ["docx"])[0].lower()
            try:
                db  = get_db()
                row = db.execute("SELECT * FROM ai_comparisons WHERE id=?", [cmp_id]).fetchone()
                if not row:
                    self._json({"error": f"comparison_id={cmp_id} 없음"}, 404)
                    return
                c = dict(row)
                # article dict 형식으로 변환
                a = {
                    "headline":    f"{c.get('corp_name','')} 비교분석",
                    "subheadline": f"{c.get('report_type_a','')} → {c.get('report_type_b','')}",
                    "corp_name":   c.get("corp_name", ""),
                    "model":       c.get("model", ""),
                    "created_at":  c.get("analyzed_at", ""),
                    "content":     c.get("result", ""),
                    "keywords":    "[]",
                    "editor_note": None,
                }
                if fmt == "docx":
                    buf = _export_docx(a)
                    fname = _safe_fname(f"비교분석_{c.get('corp_name',cmp_id)}") + ".docx"
                    self._raw(buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", fname)
                elif fmt == "pdf":
                    buf = _export_pdf(a)
                    fname = _safe_fname(f"비교분석_{c.get('corp_name',cmp_id)}") + ".pdf"
                    self._raw(buf, "application/pdf", fname)
                else:
                    self._json({"error": "format must be docx or pdf"}, 400)
            except Exception as e:
                log.error("comparison export error: %s", e, exc_info=True)
                self._json({"error": str(e)}, 500)
            return

        # ── /api/leads/{id} ───────────────────────────────────────────────
        m_lead = re.match(r"^/api/leads/(\d+)$", path)
        if m_lead:
            lead_id = int(m_lead.group(1))
            try:
                db = get_db()
                row = db.execute("""
                    SELECT sl.*, ac.result as ai_result
                    FROM story_leads sl
                    LEFT JOIN ai_comparisons ac ON sl.comparison_id = ac.id
                    WHERE sl.id=?
                """, [lead_id]).fetchone()
                if not row:
                    self._json({"error": f"lead_id={lead_id} 없음"}, 404)
                else:
                    # 해당 lead의 초안도 함께
                    drafts = db.execute(
                        "SELECT id, headline, status, model, created_at FROM article_drafts WHERE lead_id=? ORDER BY id DESC",
                        [lead_id]
                    ).fetchall()
                    d = dict(row)
                    d["drafts"] = [dict(x) for x in drafts]
                    self._json({"lead": d})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/stats/dashboard ───────────────────────────────────────────
        if path == "/api/stats/dashboard":
            try:
                db = get_db()
                def safe(sql, params=[], default=0):
                    try: return db.execute(sql, params).fetchone()[0]
                    except: return default

                self._json({
                    "reports":       safe("SELECT COUNT(*) FROM reports WHERE biz_content IS NOT NULL"),
                    "companies":     safe("SELECT COUNT(DISTINCT corp_code) FROM reports WHERE biz_content IS NOT NULL"),
                    "comparisons":   safe("SELECT COUNT(*) FROM ai_comparisons WHERE status='ok'"),
                    "supply_chain":  safe("SELECT COUNT(*) FROM supply_chain"),
                    "story_leads":   safe("SELECT COUNT(*) FROM story_leads"),
                    "leads_new":     safe("SELECT COUNT(*) FROM story_leads WHERE status='new'"),
                    "leads_high":    safe("SELECT COUNT(*) FROM story_leads WHERE severity >= 4"),
                    "articles":      safe("SELECT COUNT(*) FROM article_drafts"),
                    "report_types":  [dict(r) for r in db.execute(
                        "SELECT type_code, label, short_label, is_active FROM report_type_meta ORDER BY sort_order"
                    ).fetchall()],
                })
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/alert_rules ──────────────────────────────────────────────
        if path == "/api/alert_rules":
            try:
                db = get_db()
                rows = db.execute(
                    "SELECT * FROM alert_rules ORDER BY severity DESC, rule_code"
                ).fetchall()
                self._json({"rules": [dict(r) for r in rows]})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/ai/comparisons ──────────────────────────────────────────
        if path == "/api/ai/comparisons":
            try:
                db = get_db()
                corp_code = (qs.get("corp_code", [None])[0] or "").strip()
                type_a    = (qs.get("type_a",    [None])[0] or "").strip()
                type_b    = (qs.get("type_b",    [None])[0] or "").strip()
                model     = (qs.get("model",     [None])[0] or "").strip()
                status    = (qs.get("status",    ["ok"])[0] or "ok").strip()
                limit     = int(qs.get("limit",  [20])[0])
                offset    = int(qs.get("offset", [0])[0])
                q         = (qs.get("q",         [None])[0] or "").strip()
                sort      = (qs.get("sort",      ["recent"])[0] or "recent").strip()

                where = ["status = ?"]
                params = [status]
                if corp_code:
                    where.append("corp_code = ?")
                    params.append(corp_code)
                if type_a:
                    where.append("report_type_a = ?")
                    params.append(type_a)
                if type_b:
                    where.append("report_type_b = ?")
                    params.append(type_b)
                if model:
                    where.append("model = ?")
                    params.append(model)
                if q:
                    where.append("corp_name LIKE ?")
                    params.append(f"%{q}%")

                where_sql = "WHERE " + " AND ".join(where)

                # 정렬
                comp_order_sql = "analyzed_at DESC"  # recent and relevant both use this

                # ai_comparisons 테이블이 없으면 빈 결과 반환
                try:
                    total = db.execute(
                        f"SELECT COUNT(*) FROM ai_comparisons {where_sql}", params
                    ).fetchone()[0]
                    rows = db.execute(
                        f"SELECT id, corp_code, corp_name, report_type_a, report_type_b, "
                        f"model, char_count_a, char_count_b, status, analyzed_at, "
                        f"SUBSTR(result, 1, 500) as result_preview "
                        f"FROM ai_comparisons {where_sql} "
                        f"ORDER BY {comp_order_sql} LIMIT ? OFFSET ?",
                        params + [limit, offset]
                    ).fetchall()
                    items = [dict(r) for r in rows]
                except sqlite3.OperationalError:
                    total = 0
                    items = []

                # 전체 통계
                try:
                    stats = db.execute("""
                        SELECT report_type_a, report_type_b, model,
                               COUNT(*) as total,
                               SUM(CASE WHEN status='ok' THEN 1 ELSE 0 END) as ok_cnt,
                               SUM(CASE WHEN status='error' THEN 1 ELSE 0 END) as err_cnt
                        FROM ai_comparisons
                        GROUP BY report_type_a, report_type_b, model
                        ORDER BY report_type_a, report_type_b, model
                    """).fetchall()
                    stats_list = [dict(r) for r in stats]
                except sqlite3.OperationalError:
                    stats_list = []

                self._json({
                    "comparisons": items,
                    "total": total,
                    "offset": offset,
                    "limit": limit,
                    "stats": stats_list,
                })
            except Exception as e:
                log.exception("comparisons error")
                self._json({"error": str(e)}, 500)
            return

        # ── /api/ai/comparison_detail ─────────────────────────────────────
        if path == "/api/ai/comparison_detail":
            try:
                db = get_db()
                comp_id   = (qs.get("id",        [None])[0] or "").strip()
                corp_code = (qs.get("corp_code",  [None])[0] or "").strip()
                type_a    = (qs.get("type_a",     [None])[0] or "").strip()
                type_b    = (qs.get("type_b",     [None])[0] or "").strip()
                model     = (qs.get("model",      ["flash"])[0] or "flash").strip()

                try:
                    if comp_id:
                        row = db.execute(
                            "SELECT * FROM ai_comparisons WHERE id=?", [int(comp_id)]
                        ).fetchone()
                    elif corp_code and type_a and type_b:
                        row = db.execute("""
                            SELECT * FROM ai_comparisons
                            WHERE corp_code=? AND report_type_a=? AND report_type_b=? AND model=?
                            ORDER BY analyzed_at DESC LIMIT 1
                        """, [corp_code, type_a, type_b, model]).fetchone()
                    else:
                        self._json({"error": "id 또는 (corp_code+type_a+type_b) 필수"}, 400)
                        return

                    if not row:
                        self._json({"error": "분석 결과 없음"}, 404)
                        return
                    self._json(dict(row))
                except sqlite3.OperationalError:
                    self._json({"error": "ai_comparisons 테이블 없음 (배치 분석을 먼저 실행하세요)"}, 404)
            except Exception as e:
                log.exception("comparison_detail error")
                self._json({"error": str(e)}, 500)
            return

        # ── /api/financials ───────────────────────────────────────────────
        if path == "/api/financials":
            try:
                db = get_db()
                corp_code = (qs.get("corp_code", [None])[0] or "").strip()
                fiscal_year = (qs.get("year",    [None])[0] or "").strip()
                limit  = int(qs.get("limit",  [20])[0])
                offset = int(qs.get("offset", [0])[0])

                where, params = [], []
                if corp_code:
                    where.append("corp_code = ?"); params.append(corp_code)
                if fiscal_year:
                    where.append("fiscal_year = ?"); params.append(int(fiscal_year))

                where_sql = ("WHERE " + " AND ".join(where)) if where else ""
                try:
                    total = db.execute(
                        f"SELECT COUNT(*) FROM financials {where_sql}", params
                    ).fetchone()[0]
                    rows = db.execute(
                        f"SELECT id, corp_code, corp_name, stock_code, fiscal_year, report_type, "
                        f"consolidated, revenue, operating_income, net_income, total_assets, "
                        f"total_liabilities, total_equity, cash, debt_ratio, current_ratio, "
                        f"roe, operating_margin, net_margin, fetched_at "
                        f"FROM financials {where_sql} "
                        f"ORDER BY fiscal_year DESC, corp_code LIMIT ? OFFSET ?",
                        params + [limit, offset]
                    ).fetchall()
                    self._json({
                        "financials": [dict(r) for r in rows],
                        "total": total
                    })
                except sqlite3.OperationalError:
                    self._json({"financials": [], "total": 0})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ══════════════════════════════════════════════════════════════════
        #  공급망 그래프 API  /api/supply-chain/*
        # ══════════════════════════════════════════════════════════════════

        # ── /api/supply-chain/stats ───────────────────────────────────────
        if path == "/api/supply-chain/stats":
            try:
                from dart.supply_chain_graph import get_supply_chain_stats
                self._json(get_supply_chain_stats())
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/graph  (기업 중심 그래프) ────────────────────
        if path == "/api/supply-chain/graph":
            try:
                from dart.supply_chain_graph import get_company_graph
                corp_code = (qs.get("corp_code", [None])[0] or "").strip()
                if not corp_code:
                    self._json({"error": "corp_code 파라미터 필요"}, 400)
                    return
                depth     = int(qs.get("depth",     ["1"])[0])
                direction = (qs.get("direction",    ["both"])[0]).strip()
                relation  = (qs.get("relation",     ["all"])[0]).strip()
                self._json(get_company_graph(corp_code, depth=depth,
                                             direction=direction, relation=relation))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/hub  (파트너명 역방향 조회) ──────────────────
        if path == "/api/supply-chain/hub":
            try:
                from dart.supply_chain_graph import get_hub_detail
                partner_name = (qs.get("partner_name", [None])[0] or "").strip()
                if not partner_name:
                    self._json({"error": "partner_name 파라미터 필요"}, 400)
                    return
                relation = (qs.get("relation", ["all"])[0]).strip()
                self._json(get_hub_detail(partner_name, relation=relation))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/impact  (충격 전파 분석) ────────────────────
        if path == "/api/supply-chain/impact":
            try:
                from dart.supply_chain_graph import get_impact_analysis
                corp_code = (qs.get("corp_code", [None])[0] or "").strip()
                if not corp_code:
                    self._json({"error": "corp_code 파라미터 필요"}, 400)
                    return
                self._json(get_impact_analysis(corp_code))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/path  (두 기업 간 연결 경로 탐색) ──────────
        if path == "/api/supply-chain/path":
            try:
                from dart.supply_chain_graph import find_path
                from_code = (qs.get("from", [None])[0] or "").strip()
                to_code   = (qs.get("to",   [None])[0] or "").strip()
                if not from_code or not to_code:
                    self._json({"error": "from 및 to 파라미터 필요"}, 400)
                    return
                max_depth = int(qs.get("max_depth", ["3"])[0])
                self._json(find_path(from_code, to_code, max_depth=max_depth))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/hubs  (허브 기업 순위) ─────────────────────
        if path == "/api/supply-chain/hubs":
            try:
                from dart.supply_chain_graph import get_hub_rankings
                sector = (qs.get("sector", ["all"])[0]).strip()
                limit  = int(qs.get("limit", ["30"])[0])
                self._json(get_hub_rankings(sector=sector, limit=limit))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/changes  (보고서 기간별 공급망 변화) ─────────
        if path == "/api/supply-chain/changes":
            try:
                from dart.supply_chain_graph import get_supply_chain_changes
                corp_code = (qs.get("corp_code", [None])[0] or "").strip()
                if not corp_code:
                    self._json({"error": "corp_code 파라미터 필요"}, 400)
                    return
                from_report = (qs.get("from", [None])[0] or "").strip() or None
                to_report   = (qs.get("to",   [None])[0] or "").strip() or None
                self._json(get_supply_chain_changes(corp_code, from_report, to_report))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/risk-ranking  (집중도 리스크 순위) ─────────────
        if path == "/api/supply-chain/risk-ranking":
            try:
                from dart.supply_chain_graph import get_concentration_risk_ranking
                limit        = int((qs.get("limit",        ["50"])[0]) or 50)
                risk_level   = (qs.get("risk_level",   [None])[0] or "").strip() or None
                sort_by      = (qs.get("sort_by",      ["risk"])[0]).strip() or "risk"
                min_partners = int((qs.get("min_partners", ["5"])[0]) or 5)
                self._json(get_concentration_risk_ranking(
                    limit=limit, risk_level=risk_level,
                    sort_by=sort_by, min_partners=min_partners))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/profile  (공급망 프로파일) ─────────────────────
        if path == "/api/supply-chain/profile":
            try:
                from dart.supply_chain_graph import get_supply_chain_profile
                cc = (qs.get("corp_code", [None])[0] or "").strip()
                if not cc:
                    self._json({"error": "corp_code 필요"}, 400)
                    return
                self._json(get_supply_chain_profile(cc))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/industry-map  (산업별 공급망 지도) ─────────────
        if path == "/api/supply-chain/industry-map":
            try:
                from dart.supply_chain_graph import get_industry_map
                rtype = (qs.get("relation_type", [None])[0] or "").strip() or None
                self._json(get_industry_map(relation_type=rtype))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/daily-briefing  (일일 취재 브리핑) ──────────
        if path == "/api/supply-chain/daily-briefing":
            try:
                from dart.supply_chain_graph import get_daily_briefing
                limit = int((qs.get("limit", ["15"])[0]) or 15)
                self._json(get_daily_briefing(limit=limit))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/news/related-stocks  (뉴스 → 관련주 탐색) ─────────────────
        if path == "/api/news/related-stocks":
            try:
                from dart.news_analyzer import find_related_stocks
                # POST: body에 news_text, GET: query param
                if self.command == "POST":
                    length = int(self.headers.get("Content-Length", 0))
                    body = self.rfile.read(length).decode("utf-8", errors="replace")
                    payload = json.loads(body) if body else {}
                    news_text = payload.get("news_text", "")
                    top_n = int(payload.get("top_n", 30))
                else:
                    news_text = qs.get("q", [""])[0]
                    top_n = int((qs.get("top_n", ["30"])[0]) or 30)

                if not news_text.strip():
                    self._json({"error": "news_text 필수"}, 400)
                    return
                result = find_related_stocks(news_text, top_n=top_n)
                self._json(result)
            except Exception as e:
                import traceback
                self._json({"error": str(e), "trace": traceback.format_exc()}, 500)
            return

        # ── /api/supply-chain/article-draft  (기사 초안 자동 생성) ────────
        if path == "/api/supply-chain/article-draft":
            try:
                from dart.article_draft import generate_article_draft
                corp_code  = qs.get("corp_code", [""])[0]
                draft_type = qs.get("draft_type", ["partner_map"])[0]
                if not corp_code:
                    self._json({"error": "corp_code 필수"}, 400)
                    return
                self._json(generate_article_draft(corp_code, draft_type))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/change-alerts  (변화 알림 대시보드) ──────────
        if path == "/api/supply-chain/change-alerts":
            try:
                from dart.supply_chain_graph import get_change_alerts
                limit = int((qs.get("limit", ["30"])[0]) or 30)
                self._json(get_change_alerts(limit=limit))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/common  (교집합 종목 발굴) ─────────────────
        if path == "/api/supply-chain/common":
            try:
                from dart.supply_chain_graph import get_common_suppliers
                # partners=삼성전자&partners=현대자동차&partners=LG에너지솔루션
                partner_names = qs.get("partners", [])
                if not partner_names:
                    self._json({"error": "partners 파라미터 필요 (복수 가능)"}, 400)
                    return
                relation = (qs.get("relation", ["customer"])[0]).strip()
                self._json(get_common_suppliers(partner_names, relation=relation))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/concentration  (집중도 리스크) ────────────
        if path == "/api/supply-chain/concentration":
            try:
                from dart.supply_chain_graph import get_concentration_risk
                corp_code = (qs.get("corp_code", [None])[0] or "").strip()
                if not corp_code:
                    self._json({"error": "corp_code 파라미터 필요"}, 400)
                    return
                self._json(get_concentration_risk(corp_code))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/theme-article  (테마 기사 데이터) ─────────
        if path == "/api/supply-chain/theme-article":
            try:
                from dart.supply_chain_graph import get_theme_article_data
                corp_code    = (qs.get("corp_code",    [None])[0] or "").strip()
                article_type = (qs.get("type", ["supply_top10"])[0]).strip()
                if not corp_code:
                    self._json({"error": "corp_code 파라미터 필요"}, 400)
                    return
                self._json(get_theme_article_data(corp_code, article_type=article_type))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/supply-chain/risk-scenario  (리스크 시나리오 시뮬레이션) ─
        if path == "/api/supply-chain/risk-scenario":
            try:
                from dart.supply_chain_graph import get_risk_scenario
                corp_code = (qs.get("corp_code", [None])[0] or "").strip()
                scenario  = (qs.get("scenario",  ["disruption"])[0]).strip()
                if not corp_code:
                    self._json({"error": "corp_code 파라미터 필요"}, 400)
                    return
                self._json(get_risk_scenario(corp_code, scenario=scenario))
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── /api/dart/supply_chain  (기존 단순 조회 - 하위호환 유지) ──────
        if path == "/api/dart/supply_chain":
            try:
                db = get_db()
                corp_code  = (qs.get("corp_code",    [None])[0] or "").strip()
                partner    = (qs.get("partner",      [None])[0] or "").strip()
                rel_type   = (qs.get("relation",     [None])[0] or "").strip()
                limit      = int(qs.get("limit",  [100])[0])
                offset     = int(qs.get("offset", [0])[0])

                where = []
                params = []
                if corp_code:
                    where.append("corp_code = ?")
                    params.append(corp_code)
                if partner:
                    where.append("partner_name LIKE ?")
                    params.append(f"%{partner}%")
                if rel_type:
                    where.append("relation_type = ?")
                    params.append(rel_type)

                where_sql = ("WHERE " + " AND ".join(where)) if where else ""
                rows = db.execute(
                    f"SELECT * FROM supply_chain {where_sql} "
                    f"ORDER BY corp_name, relation_type LIMIT ? OFFSET ?",
                    params + [limit, offset]
                ).fetchall()

                total = db.execute(
                    f"SELECT COUNT(*) FROM supply_chain {where_sql}", params
                ).fetchone()[0]

                self._json({
                    "supply_chain": [dict(r) for r in rows],
                    "total": total,
                    "offset": offset,
                    "limit": limit,
                })
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── SPA 라우트 (페이지별 HTML 서빙) ─────────────────────────────
        _spa_routes = {
            "/leads":       "leads.html",
            "/comparisons": "comparisons.html",
            "/articles":    "articles.html",
        }
        if path in _spa_routes:
            spa_file = WEB_ROOT / _spa_routes[path]
            if spa_file.is_file():
                self._send(200, spa_file.read_bytes(), "text/html; charset=utf-8")
            else:
                self._send(404, f"Not Found: {_spa_routes[path]}".encode(), "text/plain")
            return

        # ── 정적 파일 서빙 ────────────────────────────────────────────────
        if path == "/" or path == "":
            path = "/dart_viewer.html"

        file_path = WEB_ROOT / path.lstrip("/")

        try:
            file_path.resolve().relative_to(WEB_ROOT.resolve())
        except ValueError:
            self._send(403, b"Forbidden", "text/plain")
            return

        if file_path.is_file():
            ext = file_path.suffix.lower()
            ct  = MIME.get(ext, "application/octet-stream")
            self._send(200, file_path.read_bytes(), ct)
        else:
            self._send(404, b"Not Found", "text/plain")

    def do_POST(self):
        parsed = urlparse(self.path)
        path = parsed.path

        # ── POST /api/ai/analyze ──────────────────────────────────────────
        if path == "/api/ai/analyze":
            # API 키 인증
            if not self._check_api_key():
                self._json({"error": "인증 실패 — X-Api-Key 헤더를 확인하세요"}, 403)
                return

            try:
                content_length = int(self.headers.get("Content-Length", 0))
                body = self.rfile.read(content_length)
                data = json.loads(body.decode("utf-8"))
            except Exception:
                self._json({"error": "요청 파싱 오류"}, 400)
                return

            model      = data.get("model", "flash")      # flash | pro | sonnet
            report_ids = data.get("report_ids", [])
            focus      = data.get("focus", "").strip()   # 선택적 분석 포커스

            if not report_ids:
                self._json({"error": "report_ids 필수"}, 400)
                return
            if len(report_ids) > 8:
                self._json({"error": "한 번에 최대 8개 보고서만 비교 가능"}, 400)
                return

            # 보고서 내용 로드 — biz_content(사업의 내용) 섹션만 사용
            db = get_db()
            reports = []
            skipped_no_biz = []
            for rid in report_ids:
                try:
                    row = db.execute(
                        "SELECT corp_name, report_name, report_type, rcept_dt, biz_content "
                        "FROM reports WHERE id = ?",
                        [int(rid)]
                    ).fetchone()
                except Exception:
                    continue
                if not row:
                    continue

                biz = (row["biz_content"] or "").strip()
                if not biz:
                    skipped_no_biz.append(row["report_name"] or str(rid))
                    continue

                # 접수일 포맷
                dt = row["rcept_dt"] or ""
                if len(dt) == 8:
                    dt = f"{dt[:4]}.{dt[4:6]}.{dt[6:]}"

                label = f"{row['corp_name']} / {row['report_name']} (접수: {dt})"
                # biz_content 최대 60,000자 (사업 내용만 사용하므로 한도 상향)
                reports.append({"label": label, "content": biz[:60000]})

            if not reports:
                msg = "사업의 내용(biz_content) 데이터가 없습니다"
                if skipped_no_biz:
                    msg += f": {', '.join(skipped_no_biz)}"
                self._json({"error": msg}, 404)
                return

            if skipped_no_biz:
                log.warning(f"사업내용 없어 제외된 보고서: {skipped_no_biz}")

            prompt = build_compare_prompt(reports, focus)

            log.info(f"AI 분석 요청: model={model}, 보고서 {len(reports)}개, 포커스='{focus}'")

            try:
                if model in ("flash", "pro"):
                    result = call_gemini(model, prompt)
                elif model == "sonnet":
                    result = call_claude(prompt)
                else:
                    self._json({"error": f"알 수 없는 모델: {model}"}, 400)
                    return

                increment_usage(model)
                usage = load_ai_usage()

                self._json({
                    "result": result,
                    "model": model,
                    "report_count": len(reports),
                    "usage": {
                        "flash":  {"used": usage.get("flash_used",  0), "limit": 1500},
                        "pro":    {"used": usage.get("pro_used",    0), "limit": 25},
                        "sonnet": {"used": usage.get("sonnet_used", 0), "limit": None},
                    }
                })
            except Exception as e:
                log.exception("AI 분석 오류")
                self._json({"error": str(e)}, 500)
            return

        # ── POST /api/comparisons/{id}/exclude ───────────────────────────────
        m_comp_excl = re.match(r"^/api/comparisons/(\d+)/exclude$", path)
        if m_comp_excl:
            comp_id = int(m_comp_excl.group(1))
            try:
                db = get_db()
                db.execute(
                    "UPDATE ai_comparisons SET status='excluded' WHERE id=?",
                    [comp_id]
                )
                db.commit()
                self._json({"ok": True, "comparison_id": comp_id, "status": "excluded"})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── POST /api/leads/{id}/draft ─────────────────────────────────────
        m_draft = re.match(r"^/api/leads/(\d+)/draft$", path)
        if m_draft:
            lead_id = int(m_draft.group(1))
            try:
                db = get_db()
                lead = db.execute("""
                    SELECT sl.*, ac.result as ai_result
                    FROM story_leads sl
                    LEFT JOIN ai_comparisons ac ON sl.comparison_id = ac.id
                    WHERE sl.id = ?
                """, [lead_id]).fetchone()

                if not lead:
                    self._json({"error": f"lead_id={lead_id} 없음"}, 404)
                    return

                # 기사 초안 생성 (인라인)
                prompt  = _build_article_prompt(lead, lead["ai_result"] or "")
                text, model_name = _call_gemini_article(prompt)
                article = _parse_article_json(text)

                if not article.get("headline"):
                    article = {
                        "headline":    lead["title"] or f"{lead['corp_name']} 관련 기사",
                        "subheadline": "",
                        "body":        text[:2000],
                        "keywords":    [],
                        "news_value":  "",
                        "caution":     "자동 파싱 실패 — 편집 필요",
                    }

                draft_id = _save_article_draft(db, lead_id, lead, article, model_name)
                draft = db.execute(
                    "SELECT * FROM article_drafts WHERE id = ?", [draft_id]
                ).fetchone()

                self._json({"ok": True, "draft": dict(draft), "model": model_name})
            except Exception as e:
                log.exception("기사 초안 생성 오류")
                self._json({"error": str(e)}, 500)
            return

        # ── POST /api/leads/{id}/check-news ───────────────────────────────
        m_check_news = re.match(r"^/api/leads/(\d+)/check-news$", path)
        if m_check_news:
            lead_id = int(m_check_news.group(1))
            try:
                import subprocess, sys as _sys
                script = str(Path(__file__).parent / "scripts" / "check_news_coverage.py")
                result = subprocess.run(
                    [_sys.executable, script, "--id", str(lead_id)],
                    capture_output=True, text=True, timeout=60, encoding="utf-8"
                )
                db = get_db()
                row = db.execute(
                    "SELECT news_status, news_checked_at, news_urls FROM story_leads WHERE id=?",
                    [lead_id]
                ).fetchone()
                if row:
                    self._json({
                        "ok": True,
                        "lead_id":        lead_id,
                        "news_status":    row["news_status"],
                        "news_checked_at":row["news_checked_at"],
                        "news_urls":      json.loads(row["news_urls"] or "[]"),
                    })
                else:
                    self._json({"error": "취재단서를 찾을 수 없습니다"}, 404)
            except subprocess.TimeoutExpired:
                self._json({"error": "뉴스 체크 타임아웃 (60초)"}, 504)
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── POST /api/leads/{id}/status ────────────────────────────────────
        m_status = re.match(r"^/api/leads/(\d+)/status$", path)
        if m_status:
            lead_id = int(m_status.group(1))
            try:
                content_length = int(self.headers.get("Content-Length", 0))
                body = self.rfile.read(content_length)
                data = json.loads(body.decode("utf-8"))
                new_status = data.get("status", "").strip()
                valid = {"new", "reviewing", "drafted", "published", "archived"}
                if new_status not in valid:
                    self._json({"error": f"유효한 status: {valid}"}, 400)
                    return
                db = get_db()
                db.execute(
                    "UPDATE story_leads SET status=?, updated_at=datetime('now','localtime') WHERE id=?",
                    [new_status, lead_id]
                )
                db.commit()
                self._json({"ok": True, "lead_id": lead_id, "status": new_status})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── POST /api/articles/{id}/status ─────────────────────────────────
        m_art_status = re.match(r"^/api/articles/(\d+)/status$", path)
        if m_art_status:
            art_id = int(m_art_status.group(1))
            try:
                content_length = int(self.headers.get("Content-Length", 0))
                body = self.rfile.read(content_length)
                data = json.loads(body.decode("utf-8"))
                new_status    = data.get("status",     "").strip()
                new_content   = data.get("content",    None)   # 본문 자동저장용
                new_headline  = data.get("headline",   None)
                new_subhead   = data.get("subheadline",None)
                db = get_db()
                if new_content is not None or new_headline is not None or new_subhead is not None:
                    # 본문/제목 업데이트 (자동저장)
                    sets, vals = [], []
                    if new_content  is not None: sets.append("content=?");    vals.append(new_content)
                    if new_headline is not None: sets.append("headline=?");   vals.append(new_headline)
                    if new_subhead  is not None: sets.append("subheadline=?");vals.append(new_subhead)
                    sets.append("updated_at=datetime('now','localtime')")
                    db.execute(
                        f"UPDATE article_drafts SET {', '.join(sets)} WHERE id=?",
                        vals + [art_id]
                    )
                    db.commit()
                    self._json({"ok": True, "article_id": art_id, "saved": True})
                    return
                valid = {"draft", "editing", "ready", "published", "rejected"}
                if new_status not in valid:
                    self._json({"error": f"유효한 status: {valid}"}, 400)
                    return
                db.execute(
                    "UPDATE article_drafts SET status=?, updated_at=datetime('now','localtime') WHERE id=?",
                    [new_status, art_id]
                )
                db.commit()
                self._json({"ok": True, "article_id": art_id, "status": new_status})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        # ── POST /api/settings ────────────────────────────────────────────
        if path == "/api/settings":
            client_ip = self.client_address[0]
            if client_ip not in ("127.0.0.1", "::1"):
                self._json({"error": "로컬에서만 접근 가능"}, 403)
                return
            try:
                content_length = int(self.headers.get("Content-Length", 0))
                body = self.rfile.read(content_length)
                data = json.loads(body.decode("utf-8"))
            except Exception:
                self._json({"error": "요청 파싱 오류"}, 400)
                return

            gemini_key    = data.get("gemini_key",    None)
            gemini_key2   = data.get("gemini_key2",   None)
            anthropic_key = data.get("anthropic_key", None)
            if gemini_key    is not None: gemini_key    = gemini_key.strip()
            if gemini_key2   is not None: gemini_key2   = gemini_key2.strip()
            if anthropic_key is not None: anthropic_key = anthropic_key.strip()
            updated = []

            def _update_env_key(file_key: str, new_val: str, env_key: str):
                """key=value 또는 # key=... 줄을 교체"""
                if not _env_path.exists():
                    return
                lines = _env_path.read_text(encoding="utf-8").splitlines()
                found = False
                new_lines = []
                for ln in lines:
                    stripped = ln.strip()
                    if (stripped.startswith(file_key + "=") or
                            stripped.startswith("# " + file_key + "=") or
                            stripped.startswith("#" + file_key + "=")):
                        new_lines.append(f"{file_key}={new_val}" if new_val else f"# {file_key}=")
                        found = True
                    else:
                        new_lines.append(ln)
                if not found and new_val:
                    new_lines.append(f"{file_key}={new_val}")
                _env_path.write_text("\n".join(new_lines) + "\n", encoding="utf-8")
                if new_val:
                    os.environ[env_key] = new_val
                elif env_key in os.environ:
                    del os.environ[env_key]

            if gemini_key is not None:
                _update_env_key("GEMINI_API_KEY", gemini_key, "GEMINI_API_KEY")
                updated.append("gemini")
                log.info(f"GEMINI_API_KEY {'설정됨' if gemini_key else '삭제됨'}")

            if gemini_key2 is not None:
                _update_env_key("GEMINI_API_KEY_2", gemini_key2, "GEMINI_API_KEY_2")
                updated.append("gemini2")
                log.info(f"GEMINI_API_KEY_2 {'설정됨' if gemini_key2 else '삭제됨'}")

            if anthropic_key is not None:
                _update_env_key("ANTHROPIC_API_KEY", anthropic_key, "ANTHROPIC_API_KEY")
                updated.append("anthropic")
                log.info(f"ANTHROPIC_API_KEY {'설정됨' if anthropic_key else '삭제됨'}")

            self._json({
                "ok": True,
                "updated": updated,
                "has_gemini":  bool(os.environ.get("GEMINI_API_KEY")),
                "has_gemini2": bool(os.environ.get("GEMINI_API_KEY_2")),
                "has_claude":  bool(os.environ.get("ANTHROPIC_API_KEY")),
            })
            return

        # ── POST /api/news/related-stocks  (뉴스 → 관련주 탐색) ─────────────
        if path == "/api/news/related-stocks":
            try:
                from dart.news_analyzer import find_related_stocks
                content_length = int(self.headers.get("Content-Length", 0))
                body = self.rfile.read(content_length).decode("utf-8", errors="replace")
                payload = json.loads(body) if body else {}
                news_text = payload.get("news_text", "").strip()
                top_n = int(payload.get("top_n", 30))
                if not news_text:
                    self._json({"error": "news_text 필수"}, 400)
                    return
                result = find_related_stocks(news_text, top_n=top_n)
                self._json(result)
            except Exception as e:
                import traceback
                self._json({"error": str(e), "trace": traceback.format_exc()}, 500)
            return

        self._json({"error": "Not found"}, 404)

    # ── PATCH 요청 처리 ────────────────────────────────────────────────────
    def do_PATCH(self):
        parsed = urlparse(self.path)
        path = parsed.path

        # ── PATCH /api/alert_rules/:id ───────────────────────────────────
        m = re.match(r"^/api/alert_rules/(\d+)$", path)
        if m:
            rule_id = int(m.group(1))
            try:
                content_length = int(self.headers.get("Content-Length", 0))
                body = self.rfile.read(content_length)
                data = json.loads(body.decode("utf-8"))
            except (json.JSONDecodeError, UnicodeDecodeError, ValueError):
                self._json({"error": "요청 파싱 오류"}, 400)
                return
            try:
                db = get_db()
                allowed = {"is_active", "title_tmpl", "description", "severity",
                           "keywords", "exclude_kw", "min_evidence_len"}
                sets, vals = [], []
                for k, v in data.items():
                    if k in allowed:
                        sets.append(f"{k}=?")
                        vals.append(v)
                if not sets:
                    self._json({"error": "변경할 필드 없음"}, 400)
                    return
                vals.append(rule_id)
                db.execute(
                    f"UPDATE alert_rules SET {', '.join(sets)} WHERE id=?", vals
                )
                db.commit()
                row = db.execute(
                    "SELECT * FROM alert_rules WHERE id=?", (rule_id,)
                ).fetchone()
                self._json({"ok": True, "rule": dict(row) if row else None})
            except Exception as e:
                self._json({"error": str(e)}, 500)
            return

        self._json({"error": "Not found"}, 404)


# ── 서버 실행 ──────────────────────────────────────────────────────────────
class ThreadingServer(socketserver.ThreadingMixIn, HTTPServer):
    daemon_threads = True


# ── 기사 내보내기 헬퍼 ──────────────────────────────────────────────────────────

def _safe_fname(s: str) -> str:
    """파일명 안전 처리 (특수문자 제거, 최대 60자)"""
    s = re.sub(r'[\\/:*?"<>|]', '_', s)
    s = s.strip().replace(' ', '_')
    return s[:60] or "article"


def _build_meta_parts(a: dict) -> list:
    """기사 dict에서 메타 정보 문자열 목록 반환 (docx/pdf 공통)"""
    parts = []
    if a.get("corp_name"):  parts.append(f"기업: {a['corp_name']}")
    if a.get("model"):      parts.append(f"AI: {a['model']}")
    if a.get("created_at"): parts.append(f"작성: {str(a['created_at'])[:16]}")
    return parts


# PDF용 한글 폰트 경로를 프로세스 시작 시 한 번만 탐색
_PDF_FONT_PATH: str | None = None
for _fp in [
    r"C:\Windows\Fonts\malgun.ttf",
    r"C:\Windows\Fonts\NanumGothic.ttf",
    r"C:\Windows\Fonts\gulim.ttc",
    "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
    "/System/Library/Fonts/AppleGothic.ttf",
]:
    if os.path.exists(_fp):
        _PDF_FONT_PATH = _fp
        break


def _export_docx(a: dict) -> bytes:
    """article dict → .docx bytes"""
    if not _HAS_DOCX:
        raise RuntimeError("python-docx 미설치. pip install python-docx")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)

    # 제목
    headline = a.get("headline") or "제목 없음"
    h = doc.add_heading(headline, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        run.font.size = Pt(20)

    # 부제목
    if a.get("subheadline"):
        sub = doc.add_paragraph(a["subheadline"])
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.size = Pt(13)
            run.italic = True

    meta_parts = _build_meta_parts(a)
    if meta_parts:
        m = doc.add_paragraph("  ·  ".join(meta_parts))
        m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in m.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # 빈 줄

    # 본문 (단락 분리)
    content = a.get("content") or ""
    for para_text in content.split("\n"):
        if para_text.strip():
            p = doc.add_paragraph(para_text)
            p.paragraph_format.space_after = Pt(6)
        else:
            doc.add_paragraph()

    # 편집 메모
    if a.get("editor_note"):
        doc.add_paragraph()
        note_para = doc.add_paragraph()
        note_para.add_run("📋 편집 메모: ").bold = True
        note_para.add_run(a["editor_note"])
        note_para.paragraph_format.left_indent = Inches(0.5)

    # 키워드
    try:
        kw = json.loads(a.get("keywords") or "[]")
    except Exception:
        kw = []
    if kw:
        doc.add_paragraph()
        kp = doc.add_paragraph()
        kp.add_run("키워드: ").bold = True
        kp.add_run(", ".join(kw))
        for run in kp.runs[1:]:
            run.font.color.rgb = RGBColor(0x5B, 0x6A, 0xF0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _export_pdf(a: dict) -> bytes:
    """article dict → .pdf bytes (fpdf2 사용, 한글 폰트 자동 탐색)"""
    if not _HAS_FPDF:
        raise RuntimeError("fpdf2 미설치. pip install fpdf2")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    if _PDF_FONT_PATH:
        pdf.add_font("KorFont", fname=_PDF_FONT_PATH)
        pdf.add_font("KorFont", style="B", fname=_PDF_FONT_PATH)
        base_font = "KorFont"
    else:
        base_font = "Arial"  # 한글 깨짐 가능 — 시스템에 한글 폰트 없음

    # 제목
    pdf.set_font(base_font, style="B", size=18)
    pdf.set_text_color(192, 57, 43)
    headline = a.get("headline") or "제목 없음"
    pdf.multi_cell(0, 10, headline, align="C")
    pdf.ln(3)

    # 부제목
    if a.get("subheadline"):
        pdf.set_font(base_font, size=12)
        pdf.set_text_color(85, 85, 85)
        pdf.multi_cell(0, 7, a["subheadline"], align="C")
        pdf.ln(3)

    meta_parts = _build_meta_parts(a)
    if meta_parts:
        pdf.set_font(base_font, size=9)
        pdf.set_text_color(136, 136, 136)
        pdf.multi_cell(0, 6, "  ·  ".join(meta_parts), align="C")

    # 구분선
    pdf.ln(4)
    pdf.set_draw_color(200, 200, 200)
    pdf.set_line_width(0.5)
    pdf.line(pdf.get_x() + 10, pdf.get_y(), pdf.get_x() + 180, pdf.get_y())
    pdf.ln(6)

    # 본문
    pdf.set_font(base_font, size=11)
    pdf.set_text_color(30, 30, 30)
    content = a.get("content") or ""
    for para_text in content.split("\n"):
        if para_text.strip():
            pdf.multi_cell(0, 7, para_text)
            pdf.ln(2)
        else:
            pdf.ln(4)

    # 키워드
    try:
        kw = json.loads(a.get("keywords") or "[]")
    except Exception:
        kw = []
    if kw:
        pdf.ln(6)
        pdf.set_font(base_font, style="B", size=9)
        pdf.set_text_color(91, 106, 240)
        pdf.multi_cell(0, 6, "키워드: " + ", ".join(kw))

    return bytes(pdf.output())


def main():
    parser = argparse.ArgumentParser(description="Company Catcher DART 뷰어 서버")
    parser.add_argument("--port", type=int, default=8888)
    parser.add_argument("--host", default="0.0.0.0")
    args = parser.parse_args()

    if not DB_PATH.exists():
        log.error(f"DB 없음: {DB_PATH}")
        sys.exit(1)

    (ROOT / "logs").mkdir(exist_ok=True)
    _ensure_api_secret_key()

    log.info(f"Company Catcher DART 서버 시작 >> http://localhost:{args.port}")
    log.info(f"DB: {DB_PATH} ({DB_PATH.stat().st_size // 1024 // 1024}MB)")
    log.info(f"Gemini API: {'설정됨' if os.environ.get('GEMINI_API_KEY') else '미설정 (.env에 GEMINI_API_KEY 추가 필요)'}")
    log.info(f"Claude API: {'설정됨' if os.environ.get('ANTHROPIC_API_KEY') else '미설정 (.env에 ANTHROPIC_API_KEY 추가 필요)'}")

    server = ThreadingServer((args.host, args.port), DartHandler)
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        log.info("서버 종료")
        server.shutdown()


if __name__ == "__main__":
    main()
